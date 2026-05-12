from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "generated" / "SpecHeal-Hackathon-Brief.docx"


BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
INK = RGBColor(0x11, 0x18, 0x27)
MUTED = RGBColor(0x55, 0x65, 0x72)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_GREEN = "E8F5E9"
PALE_AMBER = "FFF4D6"


def set_run_font(run, size: int | float | None = None, bold: bool | None = None, color: RGBColor | None = None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, col_widths: list[int]):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(col_widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in col_widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, col_widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.1):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def paragraph(doc, text: str = "", style: str | None = None, bold_prefix: str | None = None):
    p = doc.add_paragraph(style=style)
    set_paragraph_spacing(p)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True, color=INK)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, color=INK)
    else:
        r = p.add_run(text)
        set_run_font(r, color=INK)
    return p


def bullet(doc, text: str):
    p = paragraph(doc, text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    return p


def numbered(doc, text: str):
    p = paragraph(doc, text, style="List Number")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    return p


def heading(doc, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    set_paragraph_spacing(p, before=16 if level == 1 else 10, after=6)
    for run in p.runs:
        set_run_font(run, size=16 if level == 1 else 13 if level == 2 else 12, bold=True, color=BLUE if level < 3 else DARK_BLUE)
    return p


def callout(doc, label: str, body: str, fill: str = LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, before=2, after=2)
    r_label = p.add_run(label + " ")
    set_run_font(r_label, 11, True, DARK_BLUE)
    r_body = p.add_run(body)
    set_run_font(r_body, 11, False, INK)
    doc.add_paragraph()


def metadata_table(doc):
    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2300, 7060])
    rows = [
        ("Project", "SpecHeal"),
        ("Team", "Merge Kalau Berani"),
        ("Event", "Refactory Hackathon 2026, Telkom Round"),
        ("Date", "12 May 2026"),
        ("Document", "Hackathon brief generated by Codex Documents skill"),
    ]
    for i, (key, value) in enumerate(rows):
        left, right = table.rows[i].cells
        shade_cell(left, LIGHT_GRAY)
        left.text = key
        right.text = value
        for cell in (left, right):
            for p in cell.paragraphs:
                set_paragraph_spacing(p, after=2)
                for run in p.runs:
                    set_run_font(run, 10.5, key == cell.text, INK)


def scenario_table(doc):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    set_table_geometry(table, [1900, 2500, 2400, 2560])
    headers = ["Scenario", "System State", "Verdict", "Expected Output"]
    for idx, h in enumerate(headers):
        cell = table.cell(0, idx)
        shade_cell(cell, LIGHT_GRAY)
        cell.text = h
    rows = [
        ("Healthy Flow", "Original selector still works", "NO_HEAL_NEEDED", "Audit report only; no Jira issue by default"),
        ("Locator Drift", "Payment action exists under a new stable selector", "HEAL", "Apply test locator patch, rerun proof, Jira Task"),
        ("Product Bug", "Required payment action is missing or unavailable", "PRODUCT BUG", "No patch; Jira Bug with evidence"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            if idx == 2 and value == "HEAL":
                shade_cell(cells[idx], PALE_GREEN)
            elif idx == 2 and value == "PRODUCT BUG":
                shade_cell(cells[idx], PALE_AMBER)
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                set_paragraph_spacing(p, after=2, line=1.05)
                for run in p.runs:
                    set_run_font(run, 9.5, row == table.rows[0], INK)


def build():
    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("SpecHeal Hackathon Brief")
    set_run_font(r, 9, False, MUTED)

    title = doc.add_paragraph()
    set_paragraph_spacing(title, after=3)
    r = title.add_run("SpecHeal Hackathon Brief")
    set_run_font(r, 24, True, INK)

    subtitle = doc.add_paragraph()
    set_paragraph_spacing(subtitle, after=14)
    r = subtitle.add_run("AI-assisted recovery cockpit for trustworthy Playwright UI test healing")
    set_run_font(r, 12, False, MUTED)

    metadata_table(doc)

    heading(doc, "Executive Summary", 1)
    paragraph(
        doc,
        "SpecHeal is an AI-assisted recovery cockpit for Playwright UI test failures. Its narrow MVP helps QA and engineering teams decide whether a failed UI test can be safely healed or whether the product behavior is actually broken.",
    )
    callout(
        doc,
        "Core thesis:",
        "SpecHeal is not just making tests green. SpecHeal makes test recovery trustworthy.",
    )

    heading(doc, "Why This Matters", 1)
    bullet(doc, "UI automation is now core engineering infrastructure, but locator drift turns healthy product behavior into noisy CI failures.")
    bullet(doc, "Blind self-healing can produce a false green: the test passes after AI picks a replacement selector, while the product requirement remains broken.")
    bullet(doc, "SpecHeal preserves trust by requiring OpenSpec guardrails, browser validation, rerun proof, and Jira-ready workflow output.")

    heading(doc, "MVP Demo Scope", 1)
    paragraph(doc, "The demo is intentionally narrow: a seeded ShopFlow Checkout app with the user flow cart -> checkout -> pay -> Payment Success.")
    scenario_table(doc)

    heading(doc, "Recovery Loop", 1)
    numbered(doc, "Run the baseline Playwright checkout test from the dashboard.")
    numbered(doc, "Capture failure evidence: screenshot, error, cleaned DOM, visible text, and candidate selectors.")
    numbered(doc, "Load the relevant OpenSpec requirement as the behavior contract.")
    numbered(doc, "Ask live OpenAI for a structured verdict and transparent reasoning trace.")
    numbered(doc, "When verdict is HEAL, validate the candidate selector in the browser before patching.")
    numbered(doc, "Apply only a controlled Playwright test-file patch, then rerun to prove Payment Success.")
    numbered(doc, "Persist the run report in PostgreSQL and publish actionable outcomes to Jira.")

    callout(
        doc,
        "Mechanism line:",
        "AI proposes. OpenSpec guards. Browser validates. Rerun proves. Jira tracks.",
        fill=LIGHT_GRAY,
    )

    heading(doc, "Architecture Snapshot", 1)
    paragraph(doc, "Recommended stack: Next.js app router, Playwright, OpenAI structured verdict generation, PostgreSQL persistence, Jira REST publishing, and Kubernetes deployment with app and database as separate services.")
    bullet(doc, "SpecHeal App Container: dashboard, API routes, orchestrator, in-process Playwright runner, OpenAI client, Jira publisher.")
    bullet(doc, "ShopFlow Checkout Demo: logical target system served by the same app container for MVP speed.")
    bullet(doc, "PostgreSQL: run history, evidence, screenshots, AI trace, applied patch preview, Jira publish result.")
    bullet(doc, "OpenSpec files: behavior contract for ShopFlow, SpecHeal recovery, and Jira integration.")

    heading(doc, "Judge-Facing Proof Points", 1)
    bullet(doc, "Healthy Flow shows no unnecessary healing.")
    bullet(doc, "Locator Drift shows a safe patch from page.click(\"#pay-now\") to page.getByTestId(\"complete-payment\").click().")
    bullet(doc, "Product Bug shows no patch and produces a Jira Bug with evidence.")
    bullet(doc, "Live OpenAI and live Jira publishing are treated as MVP requirements, not stretch features.")
    bullet(doc, "Kubernetes and PostgreSQL demonstrate production-shaped execution, even within hackathon scope.")

    heading(doc, "Delivery Guardrails", 1)
    paragraph(doc, "SpecHeal is a proposal-and-proof system, not an automatic product-code fixer. HEAL can patch only the controlled Playwright test file in the demo runtime. Product code edits, auto-commit, auto-merge, and arbitrary website testing are outside the MVP critical path.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
